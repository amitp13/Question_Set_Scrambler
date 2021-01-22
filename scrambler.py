#!/usr/bin/env python
# coding: utf-8

# In[1]:


from docx import *
import re
import json
import random
import string

# In[ ]:

def scramble(file,outfile):
    document = Document(file)  #Change filename here
    upper_limits = []
    lower_limits = []
    limits = []
    sentences = []
    out = []
    page_one = []
    
    for idx, val in enumerate(document.paragraphs):
        sentences.append(val.text)
        if(re.search('^\d{1,2}\)', val.text)):
            ulimit = idx
            upper_limits.append(ulimit)
        elif(val.text.startswith('Answer:')):
            llimit = idx
            llimit = llimit + 1
            lower_limits.append(llimit)
        
    limits = list(zip(upper_limits, lower_limits))
   
    for l in limits:
        joined = "\n".join(sentences[l[0]:l[1]])
        out.append(joined)
        
    random.shuffle(out)
    for idx,p in enumerate(out):
        new_str = re.sub('(\d+)',str(idx+1),p)
        out[idx] = new_str
    
    for p in out:
        document.add_paragraph(p)
        document.save(outfile)

input_file = "Input file Path"
output_file = "Output File Path"
scramble(input_file,output_file)
